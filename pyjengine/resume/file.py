# You should install LibreOffice app to enable doc2pdf feature:
# Download here: https://www.libreoffice.org/donate/dl/mac-x86_64/7.5.0/en-US/LibreOffice_7.5.0_MacOS_x86-64.dmg
# Related post:  https://apple.stackexchange.com/questions/80791/command-line-tool-to-convert-doc-and-docx-files-to-pdf

from .skillmatrix import generate_skill_matrix
from .sentences import generate_resume_sentences
from .metadata import generate_meta_data

from docx import Document
from .config import RESUME_TEMPLATE_PATH, TEMP_PATH, LIBREOFFICE_PATH
import re
import os
import uuid
import shutil

def _generate_resume_file(headline, summary, positions, sentences, skill_section_headers, skill_section_contents, path):
    sentence_slot_index = 0
    position_slot_index = 0
    category_slot_index = 0
    skill_slot_index = 0
    def _interpolate_data(match_obj):
        nonlocal sentence_slot_index, position_slot_index, category_slot_index, skill_slot_index
        match = match_obj.group(1)
        if match == "headline":
            return headline
        if match == "summary":
            return summary
        if match == "position":
            position_slot_index += 1
            return positions[position_slot_index - 1]
        if match == "sentence":
            sentence_slot_index += 1
            return sentences[sentence_slot_index - 1]
        if match == "category":
            category_slot_index += 1
            return skill_section_headers[category_slot_index - 1]
        if match == "skill":
            skill_slot_index += 1
            return " • ".join(skill_section_contents[skill_slot_index - 1])
        return match
    
    document = Document(f'{RESUME_TEMPLATE_PATH}/template.docx')
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if re.match("{(.*?)}", run.text):
                run.text = re.sub("{(.*?)}", _interpolate_data, run.text)
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            if row_index >= len(skill_section_headers):
                row._element.getparent().remove(row._element)
            else:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if re.match("{(.*?)}", run.text):
                                run.text = re.sub("{(.*?)}", _interpolate_data, run.text)
    temp_file_id = str(uuid.uuid4())
    temp_docxpath = f'{TEMP_PATH}/{temp_file_id}.docx'
    temp_pdfpath = f'{TEMP_PATH}/{temp_file_id}.pdf'
    document.save(temp_docxpath)
    os.chmod(temp_docxpath, 0o777)
    os.system(f'{LIBREOFFICE_PATH}/Contents/MacOS/soffice --headless --convert-to pdf --outdir "{TEMP_PATH}" "{temp_docxpath}"')
    os.remove(temp_docxpath)
    shutil.move(temp_pdfpath, path)
    return os.path.abspath(path)

def generate_resume_file(position: str, required_skills, jd: str, path: str) -> str:
    sentences = generate_resume_sentences(position, required_skills, jd)
    ( skill_section_headers, skill_section_contents ) = generate_skill_matrix(position, required_skills)
    ( headline, summary, positions ) = generate_meta_data(position, required_skills)
    _generate_resume_file(headline, summary, positions, sentences, skill_section_headers, skill_section_contents, path)
